"""Pull content out of RFP attachments (PDF / DOCX / TXT).

Returns an ExtractionResult telling the caller whether usable text came
out (method="Text") or the document looks like a scan and the caller
should send page images to a vision model (method="Vision").
"""
import base64
import io
import logging
from dataclasses import dataclass, field

logger = logging.getLogger(__name__)

# Below this many characters from a whole PDF, assume it's a scan.
MIN_TEXT_CHARS = 50
# Cap pages sent to the vision model; keeps cost/latency bounded.
MAX_VISION_PAGES = 15
# DPI for transcription rendering. 150 is legible for body text; matches the
# existing vision-path render scale.
TRANSCRIPTION_DPI = 150


@dataclass
class ExtractionResult:
    method: str                       # "Text" or "Vision"
    text: str = ""                    # populated when method == "Text"
    page_images: list[str] = field(default_factory=list)  # base64 PNGs


def pdf_page_count(file_bytes: bytes) -> int:
    """Return the number of pages in a PDF without extracting text."""
    try:
        import fitz  # PyMuPDF - fast
        with fitz.open(stream=file_bytes, filetype="pdf") as doc:
            return doc.page_count
    except ImportError:
        from pypdf import PdfReader
        return len(PdfReader(io.BytesIO(file_bytes)).pages)


def _pdf_text_fast(file_bytes: bytes) -> str:
    """Extract embedded text with PyMuPDF (10-50x faster than pypdf on large
    documents). Falls back to pypdf if PyMuPDF isn't available."""
    try:
        import fitz  # PyMuPDF
        with fitz.open(stream=file_bytes, filetype="pdf") as doc:
            return "\n".join(page.get_text() for page in doc).strip()
    except ImportError:
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(file_bytes))
        pages = []
        for i, page in enumerate(reader.pages):
            try:
                pages.append(page.extract_text() or "")
            except Exception:
                logger.warning("Failed to extract page %s", i)
        return "\n".join(pages).strip()


def extract_text_only(file_bytes: bytes, filename: str) -> str:
    """Extract embedded text without ever rendering page images.

    Used by the RAG ingest path. Unlike extract(), this never triggers the
    vision/render path - for a scanned PDF it returns whatever little text is
    present (often empty) so the caller can skip it quickly instead of
    rendering every page to PNG (which is very slow on large document sets).
    """
    lower = filename.lower()
    if lower.endswith(".pdf"):
        return _pdf_text_fast(file_bytes)
    if lower.endswith(".docx"):
        return _from_docx(file_bytes)
    if lower.endswith((".txt", ".text", ".md")):
        return file_bytes.decode("utf-8", errors="replace")
    raise ValueError(f"Unsupported file type: {filename}")


def extract(file_bytes: bytes, filename: str) -> ExtractionResult:
    lower = filename.lower()

    if lower.endswith(".pdf"):
        return _from_pdf(file_bytes)
    if lower.endswith(".docx"):
        return ExtractionResult(method="Text", text=_from_docx(file_bytes))
    if lower.endswith((".txt", ".text", ".md")):
        return ExtractionResult(
            method="Text", text=file_bytes.decode("utf-8", errors="replace"))

    raise ValueError(f"Unsupported file type: {filename}")


def _from_pdf(file_bytes: bytes) -> ExtractionResult:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(file_bytes))
    pages = []
    for i, page in enumerate(reader.pages):
        try:
            pages.append(page.extract_text() or "")
        except Exception:
            logger.warning("Failed to extract page %s", i)
    text = "\n".join(pages).strip()

    if len(text) >= MIN_TEXT_CHARS:
        return ExtractionResult(method="Text", text=text)

    logger.info("Only %d chars extracted - treating as scanned document",
                len(text))
    return ExtractionResult(method="Vision",
                            page_images=_render_pages(file_bytes))


def _render_pages(file_bytes: bytes) -> list[str]:
    """Render PDF pages to base64 PNGs for the vision model."""
    import pypdfium2 as pdfium

    pdf = pdfium.PdfDocument(file_bytes)
    images = []
    for i in range(min(len(pdf), MAX_VISION_PAGES)):
        # 150 DPI is plenty for text legibility and keeps payloads small.
        bitmap = pdf[i].render(scale=150 / 72)
        img = bitmap.to_pil()
        buf = io.BytesIO()
        img.save(buf, format="PNG", optimize=True)
        images.append(base64.b64encode(buf.getvalue()).decode())
    if len(pdf) > MAX_VISION_PAGES:
        logger.warning("PDF has %d pages; only first %d sent to vision model",
                       len(pdf), MAX_VISION_PAGES)
    pdf.close()
    return images


def _render_all_pages(file_bytes: bytes, max_pages: int) -> list[str]:
    """Render up to max_pages PDF pages to base64 PNGs (for transcription)."""
    import pypdfium2 as pdfium

    pdf = pdfium.PdfDocument(file_bytes)
    images = []
    total = len(pdf)
    for i in range(min(total, max_pages)):
        bitmap = pdf[i].render(scale=TRANSCRIPTION_DPI / 72)
        img = bitmap.to_pil()
        buf = io.BytesIO()
        img.save(buf, format="PNG", optimize=True)
        images.append(base64.b64encode(buf.getvalue()).decode())
    if total > max_pages:
        logger.warning("PDF has %d pages; only first %d transcribed",
                       total, max_pages)
    pdf.close()
    return images


TRANSCRIBE_PROMPT = (
    "Transcribe the text on this scanned document page exactly as it "
    "appears, preserving reading order. Output plain text only: no "
    "commentary, no markdown fences, no description of the page. If a page "
    "region is illegible, omit it rather than guessing."
)


def transcribe_scanned_pdf(file_bytes: bytes, max_pages: int,
                           model: str) -> str:
    """Transcribe a scanned/image-only PDF to text via a vision model.

    One page per request, in page order, so a single failure can't lose the
    whole document. Returns the concatenated transcription. Raises ValueError
    if nothing usable was produced.
    """
    import config
    from openai import OpenAI

    pages = _render_all_pages(file_bytes, max_pages)
    if not pages:
        raise ValueError("no pages rendered for transcription")

    client = OpenAI(base_url="https://openrouter.ai/api/v1",
                    api_key=config.OPENROUTER_API_KEY)

    page_texts = []
    for i, img in enumerate(pages):
        try:
            resp = client.chat.completions.create(
                model=model,
                max_tokens=2000,
                messages=[{
                    "role": "user",
                    "content": [
                        {"type": "text", "text": TRANSCRIBE_PROMPT},
                        {"type": "image_url",
                         "image_url": {"url": f"data:image/png;base64,{img}"}},
                    ],
                }],
            )
            text = (resp.choices[0].message.content or "").strip()
        except Exception:
            logger.exception("Transcription failed on page %d", i)
            text = ""
        if text:
            page_texts.append(text)
        logger.info("Transcribed page %d/%d (%d chars)", i + 1, len(pages),
                    len(text))

    combined = "\n\n".join(page_texts).strip()
    if not combined:
        raise ValueError("transcription produced no text")
    return combined


def _from_docx(file_bytes: bytes) -> str:
    import docx

    doc = docx.Document(io.BytesIO(file_bytes))
    return "\n".join(p.text for p in doc.paragraphs).strip()
